import spacy
import re
import string
import difflib
from PyPDF2 import PdfReader
from docx import Document as DocxDocument
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
from nltk.corpus import stopwords
from nltk.stem import WordNetLemmatizer
from nltk.tokenize import word_tokenize
import nltk
from collections import Counter
import math

# Download required NLTK data
try:
    nltk.data.find('tokenizers/punkt')
except LookupError:
    nltk.download('punkt')

try:
    nltk.data.find('tokenizers/punkt_tab')
except LookupError:
    nltk.download('punkt_tab')

try:
    nltk.data.find('corpora/stopwords')
except LookupError:
    nltk.download('stopwords')

try:
    nltk.data.find('corpora/wordnet')
except LookupError:
    nltk.download('wordnet')

# Load the spaCy model
nlp = spacy.load('en_core_web_sm')

# Initialize NLTK components
stop_words = set(stopwords.words('english'))
lemmatizer = WordNetLemmatizer()

def extract_text_from_file(file_obj, file_extension):
    """
    Extract text content from uploaded file based on file type
    """
    try:
        if file_extension.lower() == '.pdf':
            return extract_text_from_pdf(file_obj)
        elif file_extension.lower() in ['.docx', '.doc']:
            return extract_text_from_docx(file_obj)
        elif file_extension.lower() == '.txt':
            return extract_text_from_txt(file_obj)
        else:
            raise ValueError(f"Unsupported file type: {file_extension}")
    except Exception as e:
        raise ValueError(f"Error processing file: {str(e)}")

def extract_text_from_pdf(file_obj):
    """
    Extract text from PDF file with improved handling
    """
    try:
        pdf_reader = PdfReader(file_obj)
        text_parts = []

        for page in pdf_reader.pages:
            page_text = page.extract_text()

            # Clean up common PDF extraction artifacts
            if page_text:
                # Remove excessive whitespace
                page_text = re.sub(r'\n\s*\n', '\n', page_text)
                page_text = re.sub(r' +', ' ', page_text)

                # Remove page numbers and common PDF artifacts
                page_text = re.sub(r'\n\d+\n', '\n', page_text)
                page_text = re.sub(r'^\d+\s*$', '', page_text, flags=re.MULTILINE)

                text_parts.append(page_text)

        full_text = '\n'.join(text_parts)

        # Final cleanup
        full_text = re.sub(r'\n{3,}', '\n\n', full_text)  # Max 2 consecutive newlines
        full_text = full_text.strip()

        return full_text

    except Exception as e:
        print(f"PDF extraction error: {e}")
        # Fallback to basic extraction
        pdf_reader = PdfReader(file_obj)
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text() + "\n"
        return text.strip()

def extract_text_from_docx(file_obj):
    """
    Extract text from DOCX file with improved handling
    """
    try:
        doc = DocxDocument(file_obj)
        text_parts = []

        for paragraph in doc.paragraphs:
            para_text = paragraph.text.strip()
            if para_text:  # Only add non-empty paragraphs
                text_parts.append(para_text)

        full_text = '\n'.join(text_parts)

        # Clean up common DOCX extraction artifacts
        full_text = re.sub(r'\n{3,}', '\n\n', full_text)  # Max 2 consecutive newlines
        full_text = re.sub(r' +', ' ', full_text)  # Normalize spaces

        return full_text.strip()

    except Exception as e:
        print(f"DOCX extraction error: {e}")
        # Fallback to basic extraction
        doc = DocxDocument(file_obj)
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        return text.strip()

def extract_text_from_txt(file_obj):
    """
    Extract text from TXT file
    """
    content = file_obj.read()
    # Try different encodings
    for encoding in ['utf-8', 'latin-1', 'cp1252']:
        try:
            return content.decode(encoding).strip()
        except UnicodeDecodeError:
            continue
    raise ValueError("Unable to decode text file with supported encodings")

def preprocess_text(text, aggressive=True):
    """
    Preprocess text for better similarity calculation
    aggressive=False preserves more formatting for identical document detection
    """
    if not text:
        return ""

    # Convert to lowercase
    text = text.lower()

    if aggressive:
        # Remove punctuation
        text = text.translate(str.maketrans('', '', string.punctuation))

        # Remove extra whitespace but preserve some structure
        text = re.sub(r'\s+', ' ', text).strip()

        # Tokenize
        tokens = word_tokenize(text)

        # Remove stopwords and lemmatize
        tokens = [lemmatizer.lemmatize(word) for word in tokens if word not in stop_words and len(word) > 2]

        return ' '.join(tokens)
    else:
        # Less aggressive preprocessing - preserve more content
        # Remove only excessive whitespace
        text = re.sub(r'\s+', ' ', text).strip()

        # Remove only problematic characters but keep most punctuation
        text = re.sub(r'[^\w\s\.\-\@\/\:]', '', text)

        return text

def calculate_jaccard_similarity(text1, text2):
    """
    Calculate Jaccard similarity between two texts
    """
    if not text1 or not text2:
        return 0.0

    # Use less aggressive preprocessing for better accuracy
    text1_processed = preprocess_text(text1, aggressive=False)
    text2_processed = preprocess_text(text2, aggressive=False)

    # Get word sets
    words1 = set(text1_processed.split())
    words2 = set(text2_processed.split())

    # Calculate Jaccard similarity
    intersection = len(words1.intersection(words2))
    union = len(words1.union(words2))

    if union == 0:
        return 0.0

    return intersection / union

def calculate_tfidf_similarity(text1, text2):
    """
    Calculate TF-IDF cosine similarity between two texts
    """
    if not text1.strip() or not text2.strip():
        return 0.0

    try:
        # Use less aggressive preprocessing for better content preservation
        text1_processed = preprocess_text(text1, aggressive=False)
        text2_processed = preprocess_text(text2, aggressive=False)

        if not text1_processed or not text2_processed:
            return 0.0

        # Create TF-IDF vectors with better parameters
        vectorizer = TfidfVectorizer(
            max_features=1000,
            ngram_range=(1, 2),
            min_df=1,  # Allow single terms
            max_df=1.0,  # Allow all terms
            sublinear_tf=True,
            stop_words='english'
        )

        # Handle case where texts are too short
        combined_text = [text1_processed, text2_processed]
        if len(combined_text[0].split()) < 3 or len(combined_text[1].split()) < 3:
            # For very short texts, use simple word overlap
            words1 = set(text1_processed.lower().split())
            words2 = set(text2_processed.lower().split())
            if not words1 or not words2:
                return 0.0
            return len(words1.intersection(words2)) / len(words1.union(words2))

        # Fit and transform
        tfidf_matrix = vectorizer.fit_transform(combined_text)

        # Calculate cosine similarity
        similarity = cosine_similarity(tfidf_matrix[0:1], tfidf_matrix[1:2])[0][0]

        return similarity

    except Exception as e:
        print(f"TF-IDF similarity error: {e}")
        # Fallback to simple word overlap
        words1 = set(text1.lower().split())
        words2 = set(text2.lower().split())
        if not words1 or not words2:
            return 0.0
        return len(words1.intersection(words2)) / len(words1.union(words2))

def calculate_semantic_similarity(text1, text2):
    """
    Calculate semantic similarity using spaCy (improved version)
    """
    if not text1.strip() or not text2.strip():
        return 0.0

    try:
        # Use larger chunks for better semantic understanding
        doc1 = nlp(text1[:10000])  # Limit text length
        doc2 = nlp(text2[:10000])

        # Calculate similarity
        similarity = doc1.similarity(doc2)

        # Adjust similarity based on content length difference
        # Penalize if texts are very different in length
        len_ratio = min(len(text1), len(text2)) / max(len(text1), len(text2))
        if len_ratio < 0.5:
            similarity *= 0.7  # Reduce similarity for very different lengths

        # Check for common words as a baseline
        words1 = set(text1.lower().split())
        words2 = set(text2.lower().split())
        word_overlap = len(words1.intersection(words2)) / len(words1.union(words2)) if words1 and words2 else 0

        # Only use semantic similarity if there's reasonable word overlap
        # This prevents high semantic scores for completely different content
        if word_overlap < 0.1:  # Less than 10% word overlap
            return word_overlap  # Return just the word overlap

        # Blend semantic similarity with word overlap for more reliable results
        # Give much more weight to word overlap for better accuracy
        combined_similarity = (similarity * 0.3) + (word_overlap * 0.7)

        # Ensure similarity is between 0 and 1
        combined_similarity = max(0.0, min(1.0, combined_similarity))

        return combined_similarity

    except Exception as e:
        print(f"Semantic similarity error: {e}")
        # Fallback to word overlap
        words1 = set(text1.lower().split())
        words2 = set(text2.lower().split())
        if not words1 or not words2:
            return 0.0
        return len(words1.intersection(words2)) / len(words1.union(words2))

def calculate_direct_similarity(text1, text2):
    """
    Calculate direct text similarity for identical document detection
    Only returns high similarity for truly identical or nearly identical content
    """
    if not text1.strip() or not text2.strip():
        return 0.0

    # Normalize texts for comparison
    text1_norm = re.sub(r'\s+', ' ', text1.strip().lower())
    text2_norm = re.sub(r'\s+', ' ', text2.strip().lower())

    # If texts are identical after normalization, return 100%
    if text1_norm == text2_norm:
        return 1.0

    # For near-identical content, use difflib for better character-level comparison
    if len(text1_norm) > 10 and len(text2_norm) > 10:
        matcher = difflib.SequenceMatcher(None, text1_norm, text2_norm)
        char_similarity = matcher.ratio()

        # Only consider it highly similar if character similarity is > 95%
        # and the texts are reasonably similar in length
        len_ratio = min(len(text1_norm), len(text2_norm)) / max(len(text1_norm), len(text2_norm))

        if char_similarity > 0.95 and len_ratio > 0.8:
            return char_similarity
        elif char_similarity > 0.90 and len_ratio > 0.9:
            return char_similarity * 0.8  # Reduce for slightly different content

    # For different content, return a low similarity based on word overlap
    words1 = set(text1_norm.split())
    words2 = set(text2_norm.split())

    if not words1 or not words2:
        return 0.0

    word_overlap = len(words1.intersection(words2)) / len(words1.union(words2))

    # Return word overlap but cap it at 30% for direct similarity
    # This prevents different documents from getting high direct similarity
    return min(word_overlap, 0.3)

def calculate_keyword_overlap(text1, text2):
    """
    Calculate keyword overlap similarity
    """
    if not text1 or not text2:
        return 0.0

    # Extract important keywords (nouns, proper nouns, technical terms)
    doc1 = nlp(text1[:5000])
    doc2 = nlp(text2[:5000])

    # Get keywords (nouns and proper nouns)
    keywords1 = set()
    keywords2 = set()

    for token in doc1:
        if token.pos_ in ['NOUN', 'PROPN'] and len(token.text) > 2:
            keywords1.add(token.lemma_.lower())

    for token in doc2:
        if token.pos_ in ['NOUN', 'PROPN'] and len(token.text) > 2:
            keywords2.add(token.lemma_.lower())

    if not keywords1 or not keywords2:
        return 0.0

    # Calculate overlap
    intersection = len(keywords1.intersection(keywords2))
    union = len(keywords1.union(keywords2))

    return intersection / union if union > 0 else 0.0
    """
    Calculate keyword overlap similarity
    """
    if not text1 or not text2:
        return 0.0

    # Extract important keywords (nouns, proper nouns, technical terms)
    doc1 = nlp(text1[:5000])
    doc2 = nlp(text2[:5000])

    # Get keywords (nouns and proper nouns)
    keywords1 = set()
    keywords2 = set()

    for token in doc1:
        if token.pos_ in ['NOUN', 'PROPN'] and len(token.text) > 2:
            keywords1.add(token.lemma_.lower())

    for token in doc2:
        if token.pos_ in ['NOUN', 'PROPN'] and len(token.text) > 2:
            keywords2.add(token.lemma_.lower())

    if not keywords1 or not keywords2:
        return 0.0

    # Calculate overlap
    intersection = len(keywords1.intersection(keywords2))
    union = len(keywords1.union(keywords2))

    return intersection / union if union > 0 else 0.0

def normalize_extracted_text(text):
    """
    Normalize text extracted from different formats for better comparison
    """
    if not text:
        return ""

    # Convert to lowercase for case-insensitive comparison
    text = text.lower()

    # Normalize whitespace
    text = re.sub(r'\s+', ' ', text)

    # Remove common formatting artifacts
    text = re.sub(r'[^\w\s\.\-\@\/\:]', '', text)  # Keep basic punctuation

    # Remove extra whitespace around punctuation
    text = re.sub(r'\s*([\.!\?])\s*', r'\1 ', text)

    # Normalize common abbreviations and terms
    text = re.sub(r'\bdr\.\s*', 'doctor ', text)
    text = re.sub(r'\bco\.\s*', 'company ', text)
    text = re.sub(r'\binc\.\s*', 'incorporated ', text)
    text = re.sub(r'\bcorp\.\s*', 'corporation ', text)

    return text.strip()

def calculate_similarity(text1, text2):
    """
    Calculate comprehensive similarity percentage between two texts
    using multiple algorithms and weighted combination
    """
    if not text1.strip() or not text2.strip():
        return 0.0

    try:
        # Create normalized versions for better cross-format comparison
        text1_norm = normalize_extracted_text(text1)
        text2_norm = normalize_extracted_text(text2)

        # Calculate different similarity metrics using normalized text
        jaccard_sim = calculate_jaccard_similarity(text1_norm, text2_norm)
        tfidf_sim = calculate_tfidf_similarity(text1_norm, text2_norm)
        semantic_sim = calculate_semantic_similarity(text1, text2)  # Use original for context
        keyword_sim = calculate_keyword_overlap(text1, text2)
        direct_sim = calculate_direct_similarity(text1_norm, text2_norm)

        # If direct similarity is very high, boost the result significantly
        if direct_sim > 0.9:
            return 95.0 + (direct_sim - 0.9) * 50  # Can reach up to 100%

        # Weighted combination (adjust weights based on document type)
        # For resumes/technical documents, TF-IDF and keyword overlap are more important
        weights = {
            'tfidf': 0.35,     # Most important for technical content
            'keyword': 0.35,   # Important for skill matching
            'jaccard': 0.25,   # Good for general overlap
            'semantic': 0.03,  # Much less weight due to unreliability
            'direct': 0.02     # Small boost for identical content
        }

        combined_similarity = (
            weights['tfidf'] * tfidf_sim +
            weights['keyword'] * keyword_sim +
            weights['jaccard'] * jaccard_sim +
            weights['semantic'] * semantic_sim +
            weights['direct'] * direct_sim
        )

        # Apply sigmoid function to make results more reasonable
        # This prevents extreme values and provides better distribution
        if combined_similarity > 0:
            combined_similarity = 1 / (1 + math.exp(-5 * (combined_similarity - 0.5)))

        # Convert to percentage and round to 2 decimal places
        similarity_percentage = round(combined_similarity * 100, 2)

        # Ensure reasonable bounds
        similarity_percentage = max(0.0, min(100.0, similarity_percentage))

        return similarity_percentage

    except Exception as e:
        print(f"Similarity calculation error: {e}")
        # Fallback to direct similarity if other methods fail
        direct_sim = calculate_direct_similarity(text1, text2)
        return round(direct_sim * 100, 2)

def get_file_extension(filename):
    """
    Get file extension from filename
    """
    return '.' + filename.split('.')[-1] if '.' in filename else ''